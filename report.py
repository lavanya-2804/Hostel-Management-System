import mysql.connector
from docx import Document

# Establish database connection
db = mysql.connector.connect(
    host="localhost",
    user="root",
    password="yoursqlpassword",
    database="hostel"
)

# Create cursor
cursor = db.cursor()

# Function to generate the report
def generate_report():
    # Create a new Word document
    document = Document()

    # Add a title to the document
    document.add_heading("Hostel Management System Report", level=1)

    # Retrieve student details
    cursor.execute("""
        SELECT students.id, students.name, students.roll_no, students.department, rooms.room_no, students.fees_paid, students.fees_due
        FROM students
        LEFT JOIN rooms ON students.room_no = rooms.room_no
    """)
    student_data = cursor.fetchall()

    # Add the Student Details section to the document
    document.add_heading("Student Details", level=2)
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table_header = table.rows[0].cells
    table_header[0].text = "ID"
    table_header[1].text = "Name"
    table_header[2].text = "Roll No"
    table_header[3].text = "Department"
    table_header[4].text = "Room No"
    table_header[5].text = "Fees Paid"
    table_header[6].text = "Fees Due"
    for student in student_data:
        row = table.add_row().cells
        row[0].text = str(student[0])
        row[1].text = student[1]
        row[2].text = student[2]
        row[3].text = student[3]
        row[4].text = str(student[4]) if student[4] else "Not Assigned"
        row[5].text = "{:.2f}".format(student[5])
        row[6].text = "{:.2f}".format(student[6])

    # Retrieve room occupancy
    cursor.execute("""
        SELECT room_no, is_occupied
        FROM rooms
    """)
    room_data = cursor.fetchall()

    # Add the Room Occupancy section to the document
    document.add_heading("Room Occupancy", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table_header = table.rows[0].cells
    table_header[0].text = "Room No"
    table_header[1].text = "Occupancy"
    for room in room_data:
        row = table.add_row().cells
        row[0].text = str(room[0])
        occupancy = "Occupied" if room[1] else "Vacant"
        row[1].text = occupancy

    # Calculate total fees collection and dues
    cursor.execute("""
        SELECT SUM(fees_paid), SUM(fees_due)
        FROM students
    """)
    result = cursor.fetchone()
    total_fees_paid = result[0] or 0.0
    total_fees_due = result[1] or 0.0

    # Add the Fees Status section to the document
    document.add_heading("Fees Status", level=2)
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Total Fees Paid"
    table.cell(0, 1).text = "{:.2f}".format(total_fees_paid)
    table.cell(1, 0).text = "Total Fees Due"
    table.cell(1, 1).text = "{:.2f}".format(total_fees_due)

    # Save the document
    document.save("hostel_management_report.docx")

# Generate the report
generate_report()

# Close the database connection
db.close()

print("Report generated successfully.")
